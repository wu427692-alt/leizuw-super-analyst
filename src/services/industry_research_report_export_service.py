# -*- coding: utf-8 -*-
"""Publication-grade Word/PDF exports for industry and company research.

The web workbench contains task state, QA and ingestion diagnostics.  Those
fields are deliberately excluded here: an exported artifact is a research
report, not a worker log.  Both formats share the same citation numbering and
report-only structure.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from html import escape as html_escape
from io import BytesIO
import re
from tempfile import TemporaryDirectory
from pathlib import Path
from typing import Any, Dict, Iterable, List, Mapping, Sequence, Tuple


_CITATION_RE = re.compile(
    r"\[((?:report|note|event|announcement|filing|financial|industry-peer|web|audio):[^\]\s]+)\]",
    re.IGNORECASE,
)
_FIGURE_RE = re.compile(r"(?:图表)?【([^｜】]+)(?:｜([^】]+))?】")
_MARKDOWN_PREFIX_RE = re.compile(r"^\s{0,3}(#{1,6})\s+(.+?)\s*$")
_LIST_RE = re.compile(r"^\s*(?:[-*+]\s+|\d+[.)]\s+)(.+)$")


@dataclass(frozen=True)
class ResearchReportArtifact:
    content: bytes
    media_type: str
    extension: str


class IndustryResearchReportExportService:
    """Render a completed project as a clean, reader-facing report."""

    WORD_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    PDF_MEDIA_TYPE = "application/pdf"
    ACCENT = "9B4B1B"
    ACCENT_RGB = (155, 75, 27)
    GOLD_RGB = (201, 159, 85)
    BLUE_GREY_RGB = (83, 126, 142)
    TABLE_FILL = "F4EBDD"

    def render(self, project: Mapping[str, Any], output_format: str) -> ResearchReportArtifact:
        fmt = str(output_format or "").lower().strip()
        if fmt not in {"docx", "pdf"}:
            raise ValueError("只支持 Word 或 PDF 报告")
        prepared = self._prepare(project)
        if fmt == "docx":
            return ResearchReportArtifact(self._render_docx(prepared), self.WORD_MEDIA_TYPE, "docx")
        return ResearchReportArtifact(self._render_pdf(prepared), self.PDF_MEDIA_TYPE, "pdf")

    @classmethod
    def _prepare(cls, project: Mapping[str, Any]) -> Dict[str, Any]:
        report = project.get("report") if isinstance(project.get("report"), Mapping) else {}
        if not report:
            raise ValueError("报告尚未生成完成")
        chapters = [item for item in report.get("chapters") or [] if isinstance(item, Mapping)]
        long_form = str(report.get("long_form_report") or "").strip()
        if not chapters and not long_form:
            raise ValueError("长篇正文仍在后台生成")
        snapshot = project.get("snapshot") if isinstance(project.get("snapshot"), Mapping) else {}
        evidence = [item for item in snapshot.get("evidence") or [] if isinstance(item, Mapping)]
        evidence_by_id = {
            str(item.get("evidence_id") or item.get("evidenceId") or ""): dict(item)
            for item in evidence
            if item.get("evidence_id") or item.get("evidenceId")
        }
        bodies = [
            str(report.get("one_sentence") or report.get("oneSentence") or ""),
            str(report.get("executive_summary") or report.get("executiveSummary") or ""),
            *[str(item.get("body_markdown") or item.get("bodyMarkdown") or "") for item in chapters],
        ]
        if not bodies and long_form:
            bodies = [long_form]
        citation_ids: List[str] = []
        for body in bodies:
            for evidence_id in _CITATION_RE.findall(body):
                if evidence_id not in citation_ids:
                    citation_ids.append(evidence_id)
        citation_numbers = {evidence_id: index + 1 for index, evidence_id in enumerate(citation_ids)}
        subject = report.get("subject") if isinstance(report.get("subject"), Mapping) else {}
        if not subject and isinstance(snapshot.get("subject"), Mapping):
            subject = snapshot.get("subject")
        completed_at = report.get("research_cutoff") or report.get("researchCutoff") or project.get("completed_at") or project.get("completedAt")
        return {
            "topic": str(project.get("topic") or "深度研究"),
            "research_type": str(project.get("research_type") or project.get("researchType") or report.get("research_type") or report.get("researchType") or "industry"),
            "symbol": str((subject or {}).get("symbol") or ""),
            "date": cls._display_date(completed_at),
            "thesis": str(report.get("one_sentence") or report.get("oneSentence") or "").strip(),
            "summary": str(report.get("executive_summary") or report.get("executiveSummary") or "").strip(),
            "chapters": chapters,
            "long_form": long_form,
            "figures": [item for item in report.get("visualizations") or [] if isinstance(item, Mapping)],
            "citations": citation_numbers,
            "sources": [
                {"number": citation_numbers[evidence_id], **evidence_by_id.get(evidence_id, {"evidence_id": evidence_id})}
                for evidence_id in citation_ids
            ],
        }

    @staticmethod
    def _display_date(value: Any) -> str:
        raw = str(value or "").strip()
        if not raw:
            return datetime.now().strftime("%Y年%m月%d日")
        try:
            parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
            return parsed.strftime("%Y年%m月%d日")
        except ValueError:
            return raw[:10].replace("-", "年", 1).replace("-", "月", 1) + ("日" if len(raw) >= 10 else "")

    @classmethod
    def _clean_text(cls, value: Any, citations: Mapping[str, int]) -> str:
        text = str(value or "")
        text = _CITATION_RE.sub(lambda match: f"@@CITE:{citations.get(match.group(1), 0)}@@" if citations.get(match.group(1)) else "", text)
        text = _FIGURE_RE.sub(lambda match: f"见图表 {match.group(1)}" + (f" {match.group(2)}" if match.group(2) else ""), text)
        return text.replace("```markdown", "").replace("```", "").strip()

    @staticmethod
    def _markdown_blocks(markdown: str) -> Iterable[Tuple[str, Any]]:
        lines = str(markdown or "").replace("\r\n", "\n").split("\n")
        index = 0
        while index < len(lines):
            line = lines[index].rstrip()
            if not line.strip():
                index += 1
                continue
            heading = _MARKDOWN_PREFIX_RE.match(line)
            if heading:
                yield "heading", (len(heading.group(1)), heading.group(2).strip())
                index += 1
                continue
            if "|" in line and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1]):
                table_lines = [line]
                index += 2
                while index < len(lines) and "|" in lines[index] and lines[index].strip():
                    table_lines.append(lines[index].rstrip())
                    index += 1
                rows = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in table_lines]
                yield "table", rows
                continue
            list_match = _LIST_RE.match(line)
            if list_match:
                items: List[str] = []
                while index < len(lines):
                    match = _LIST_RE.match(lines[index].rstrip())
                    if not match:
                        break
                    items.append(match.group(1).strip())
                    index += 1
                yield "list", items
                continue
            if line.lstrip().startswith(">"):
                quote: List[str] = []
                while index < len(lines) and lines[index].lstrip().startswith(">"):
                    quote.append(lines[index].lstrip()[1:].strip())
                    index += 1
                yield "quote", " ".join(quote)
                continue
            paragraph = [line.strip()]
            index += 1
            while index < len(lines):
                next_line = lines[index].rstrip()
                if not next_line.strip() or _MARKDOWN_PREFIX_RE.match(next_line) or _LIST_RE.match(next_line) or next_line.lstrip().startswith(">"):
                    break
                if "|" in next_line and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-{3,}", lines[index + 1]):
                    break
                paragraph.append(next_line.strip())
                index += 1
            yield "paragraph", " ".join(paragraph)

    @classmethod
    def _render_docx(cls, prepared: Mapping[str, Any]) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.5)
        section.top_margin, section.bottom_margin = Cm(2.2), Cm(2.2)
        section.header_distance, section.footer_distance = Cm(1.0), Cm(1.0)

        styles = document.styles
        if Path("/System/Library/Fonts/Hiragino Sans GB.ttc").is_file():
            body_font = heading_font = "Hiragino Sans GB"
        elif Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf").is_file():
            body_font = heading_font = "Arial Unicode MS"
        elif Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc").is_file():
            body_font = heading_font = "Noto Sans CJK SC"
        else:
            body_font, heading_font = "宋体", "微软雅黑"
        normal = styles["Normal"]
        normal.font.name, normal.font.size = body_font, Pt(10.5)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
        normal.paragraph_format.line_spacing = 1.55
        normal.paragraph_format.space_after = Pt(7)
        for name, size, color in (("Title", 30, cls.ACCENT), ("Heading 1", 18, cls.ACCENT), ("Heading 2", 14, "222222"), ("Heading 3", 12, "333333")):
            style = styles[name]
            style.font.name, style.font.size, style.font.bold = heading_font, Pt(size), True
            style._element.rPr.rFonts.set(qn("w:eastAsia"), heading_font)
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(14)
            style.paragraph_format.space_after = Pt(8)

        header = section.header.paragraphs[0]
        header.text = f"{prepared['topic']}深度研究报告"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.size, run.font.color.rgb = Pt(8.5), RGBColor(130, 130, 130)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run()
        fld_char = OxmlElement("w:fldChar"); fld_char.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_char, instr, fld_end])

        for _ in range(4):
            document.add_paragraph()
        symbol = str(prepared.get("symbol") or "")
        if symbol:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_run = paragraph.add_run(symbol)
            cover_run.bold, cover_run.font.size = True, Pt(16)
            cover_run.font.color.rgb = RGBColor.from_string(cls.ACCENT)
        title = document.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(f"{prepared['topic']}\n深度研究报告")
        if prepared.get("thesis"):
            thesis = document.add_paragraph()
            thesis.alignment = WD_ALIGN_PARAGRAPH.CENTER
            thesis.paragraph_format.space_before = Pt(20)
            thesis.paragraph_format.left_indent = thesis.paragraph_format.right_indent = Cm(1.2)
            thesis_run = thesis.add_run(str(prepared["thesis"]))
            thesis_run.font.size, thesis_run.font.color.rgb = Pt(13), RGBColor(70, 70, 70)
        date = document.add_paragraph(str(prepared["date"]))
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date.paragraph_format.space_before = Pt(30)
        document.add_page_break()

        document.add_heading("核心结论", level=1)
        cls._docx_rich_paragraph(document, prepared.get("summary") or prepared.get("thesis") or "", prepared["citations"], lead=True)
        chapters = list(prepared.get("chapters") or [])
        document.add_heading("目录", level=1)
        if chapters:
            for index, chapter in enumerate(chapters, 1):
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{index}. ").bold = True
                paragraph.add_run(str(chapter.get("title") or f"第{index}章"))
        document.add_page_break()

        if chapters:
            for index, chapter in enumerate(chapters, 1):
                title_text = re.sub(r"^第?\s*\d+[章节\.、\s]*", "", str(chapter.get("title") or "")).strip()
                document.add_heading(f"{index}. {title_text or f'第{index}章'}", level=1)
                summary = str(chapter.get("summary") or "").strip()
                if summary:
                    cls._docx_rich_paragraph(document, summary, prepared["citations"], lead=True)
                body = cls._clean_text(chapter.get("body_markdown") or chapter.get("bodyMarkdown") or "", prepared["citations"])
                cls._append_docx_markdown(document, body, prepared["citations"])
        else:
            cls._append_docx_markdown(document, cls._clean_text(prepared.get("long_form"), prepared["citations"]), prepared["citations"])

        figures = list(prepared.get("figures") or [])
        if figures:
            document.add_heading("关键图表", level=1)
            with TemporaryDirectory(prefix="industry-report-figures-") as directory:
                for index, figure in enumerate(figures, 1):
                    document.add_heading(f"图表 {index}  {figure.get('title') or '数据图表'}", level=2)
                    path = Path(directory) / f"figure-{index}.png"
                    cls._render_figure_png(figure, path)
                    document.add_picture(str(path), width=Inches(6.1))
                    caption = document.add_paragraph(f"资料来源：{figure.get('source') or '本次研究固定证据快照'}")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for caption_run in caption.runs:
                        caption_run.font.size, caption_run.font.color.rgb = Pt(8), RGBColor(145, 145, 145)

        if prepared.get("sources"):
            document.add_heading("资料来源", level=1)
            for source in prepared["sources"]:
                number = source.get("number")
                source_name = source.get("source") or source.get("source_name") or source.get("sourceName") or source.get("kind") or "资料来源"
                title_text = source.get("title") or source.get("summary") or source.get("evidence_id") or source.get("evidenceId")
                date_text = source.get("date") or ""
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Cm(-0.55)
                paragraph.paragraph_format.left_indent = Cm(0.55)
                paragraph.add_run(f"{number}. ").bold = True
                paragraph.add_run(" · ".join(str(value) for value in (source_name, date_text, title_text) if value))
                if source.get("url"):
                    url_run = paragraph.add_run(f"\n{source['url']}")
                    url_run.font.size, url_run.font.color.rgb = Pt(8), RGBColor(90, 110, 135)

        document.add_paragraph()
        disclaimer = document.add_paragraph("本报告基于报告所列资料整理，仅供研究参考，不构成投资建议。")
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for disclaimer_run in disclaimer.runs:
            disclaimer_run.font.size, disclaimer_run.font.color.rgb = Pt(8), RGBColor(140, 140, 140)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @classmethod
    def _append_docx_markdown(cls, document: Any, markdown: str, citations: Mapping[str, int]) -> None:
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import RGBColor, Pt

        for kind, value in cls._markdown_blocks(markdown):
            if kind == "heading":
                level, text = value
                document.add_heading(re.sub(r"^\d+(?:\.\d+)*[\.、\s]+", "", text), level=min(max(level, 2), 3))
            elif kind == "table":
                rows = value
                if not rows:
                    continue
                width = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=width)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Table Grid"
                for row_index, row in enumerate(rows):
                    for column_index in range(width):
                        cell = table.cell(row_index, column_index)
                        cell.text = row[column_index] if column_index < len(row) else ""
                        if row_index == 0:
                            shading = OxmlElement("w:shd"); shading.set(qn("w:fill"), cls.TABLE_FILL); cell._tc.get_or_add_tcPr().append(shading)
                            for run in cell.paragraphs[0].runs:
                                run.bold, run.font.color.rgb = True, RGBColor.from_string(cls.ACCENT)
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_after = Pt(2)
                document.add_paragraph()
            elif kind == "list":
                for item in value:
                    paragraph = document.add_paragraph(style="List Bullet")
                    cls._docx_add_runs(paragraph, item, citations)
            elif kind == "quote":
                paragraph = document.add_paragraph()
                paragraph.style = document.styles["Quote"]
                cls._docx_add_runs(paragraph, value, citations)
            else:
                cls._docx_rich_paragraph(document, value, citations)

    @classmethod
    def _docx_rich_paragraph(cls, document: Any, text: Any, citations: Mapping[str, int], *, lead: bool = False) -> Any:
        from docx.shared import Cm, Pt, RGBColor
        paragraph = document.add_paragraph()
        if lead:
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.right_indent = Cm(0.5)
            paragraph.paragraph_format.space_after = Pt(14)
        cls._docx_add_runs(paragraph, cls._clean_text(text, citations), citations, lead=lead)
        return paragraph

    @staticmethod
    def _docx_add_runs(paragraph: Any, text: str, citations: Mapping[str, int], *, lead: bool = False) -> None:
        from docx.shared import Pt, RGBColor
        parts = re.split(r"(@@CITE:\d+@@|\*\*.+?\*\*)", str(text or ""))
        for part in parts:
            if not part:
                continue
            citation = re.fullmatch(r"@@CITE:(\d+)@@", part)
            if citation:
                run = paragraph.add_run(citation.group(1))
                run.font.superscript, run.font.size = True, Pt(7)
                run.font.color.rgb = RGBColor.from_string(IndustryResearchReportExportService.ACCENT)
                continue
            bold = part.startswith("**") and part.endswith("**")
            run = paragraph.add_run(part[2:-2] if bold else part)
            run.bold = bold or lead
            if lead:
                run.font.size, run.font.color.rgb = Pt(12), RGBColor(65, 65, 65)

    @classmethod
    def _render_pdf(cls, prepared: Mapping[str, Any]) -> bytes:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

        font_path = cls._find_cjk_font()
        pdf_font = "LeziwuCJK" if font_path else "STSong-Light"
        if font_path:
            pdfmetrics.registerFont(TTFont(pdf_font, font_path, subfontIndex=0))
        else:
            pdfmetrics.registerFont(UnicodeCIDFont(pdf_font))
        pdfmetrics.registerFontFamily(
            pdf_font,
            normal=pdf_font,
            bold=pdf_font,
            italic=pdf_font,
            boldItalic=pdf_font,
        )
        buffer = BytesIO()
        document = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2.25 * cm, rightMargin=2.25 * cm, topMargin=2.2 * cm, bottomMargin=2.1 * cm, title=f"{prepared['topic']}深度研究报告")
        styles = getSampleStyleSheet()
        body = ParagraphStyle("CNBody", parent=styles["BodyText"], fontName=pdf_font, fontSize=10.5, leading=18, textColor=colors.HexColor("#333333"), spaceAfter=8)
        lead = ParagraphStyle("CNLead", parent=body, fontSize=12, leading=21, textColor=colors.HexColor("#3D3D3D"), leftIndent=12, rightIndent=12, spaceAfter=16)
        h1 = ParagraphStyle("CNH1", parent=styles["Heading1"], fontName=pdf_font, fontSize=19, leading=26, textColor=colors.HexColor(f"#{cls.ACCENT}"), spaceBefore=14, spaceAfter=9)
        h2 = ParagraphStyle("CNH2", parent=styles["Heading2"], fontName=pdf_font, fontSize=14, leading=20, textColor=colors.HexColor("#222222"), spaceBefore=12, spaceAfter=7)
        h3 = ParagraphStyle("CNH3", parent=h2, fontSize=11.5, leading=17)
        small = ParagraphStyle("CNSmall", parent=body, fontSize=7.5, leading=11, textColor=colors.HexColor("#8A8A8A"))
        cover_title = ParagraphStyle("CoverTitle", parent=h1, alignment=TA_CENTER, fontSize=30, leading=42, spaceBefore=34, spaceAfter=26)
        cover_symbol = ParagraphStyle("CoverSymbol", parent=h2, alignment=TA_CENTER, textColor=colors.HexColor(f"#{cls.ACCENT}"), fontSize=14)
        story: List[Any] = [Spacer(1, 4.2 * cm)]
        if prepared.get("symbol"):
            story.append(Paragraph(html_escape(str(prepared["symbol"])), cover_symbol))
        story.append(Paragraph(f"{html_escape(str(prepared['topic']))}<br/>深度研究报告", cover_title))
        if prepared.get("thesis"):
            story.append(Paragraph(cls._pdf_text(prepared["thesis"], prepared["citations"]), ParagraphStyle("CoverThesis", parent=lead, alignment=TA_CENTER)))
        story.extend([Spacer(1, 1.4 * cm), Paragraph(html_escape(str(prepared["date"])), ParagraphStyle("CoverDate", parent=body, alignment=TA_CENTER)), PageBreak()])
        story.extend([Paragraph("核心结论", h1), Paragraph(cls._pdf_text(prepared.get("summary") or prepared.get("thesis") or "", prepared["citations"]), lead), Paragraph("目录", h1)])
        chapters = list(prepared.get("chapters") or [])
        for index, chapter in enumerate(chapters, 1):
            story.append(Paragraph(html_escape(f"{index}. {chapter.get('title') or f'第{index}章'}"), body))
        story.append(PageBreak())
        if chapters:
            for index, chapter in enumerate(chapters, 1):
                title_text = re.sub(r"^第?\s*\d+[章节\.、\s]*", "", str(chapter.get("title") or "")).strip()
                story.append(Paragraph(html_escape(f"{index}. {title_text or f'第{index}章'}"), h1))
                if chapter.get("summary"):
                    story.append(Paragraph(cls._pdf_text(chapter.get("summary"), prepared["citations"]), lead))
                markdown = cls._clean_text(chapter.get("body_markdown") or chapter.get("bodyMarkdown") or "", prepared["citations"])
                cls._append_pdf_markdown(story, markdown, prepared["citations"], body, h2, h3, Table, TableStyle, colors)
        else:
            cls._append_pdf_markdown(story, cls._clean_text(prepared.get("long_form"), prepared["citations"]), prepared["citations"], body, h2, h3, Table, TableStyle, colors)

        figures = list(prepared.get("figures") or [])
        if figures:
            story.append(Paragraph("关键图表", h1))
            with TemporaryDirectory(prefix="industry-report-figures-") as directory:
                for index, figure in enumerate(figures, 1):
                    path = Path(directory) / f"figure-{index}.png"
                    cls._render_figure_png(figure, path)
                    # ReportLab reads images during ``build``. Keep the image
                    # bytes alive after this temporary directory is removed,
                    # and keep each heading with its chart and source note.
                    story.append(KeepTogether([
                        Paragraph(html_escape(f"图表 {index}  {figure.get('title') or '数据图表'}"), h2),
                        Image(BytesIO(path.read_bytes()), width=15.6 * cm, height=8.5 * cm),
                        Paragraph(html_escape(f"资料来源：{figure.get('source') or '本次研究固定证据快照'}"), small),
                    ]))

        if prepared.get("sources"):
            story.append(Paragraph("资料来源", h1))
            for source in prepared["sources"]:
                values = [source.get("source") or source.get("source_name") or source.get("sourceName") or source.get("kind") or "资料来源", source.get("date"), source.get("title") or source.get("summary") or source.get("evidence_id") or source.get("evidenceId")]
                source_text = " · ".join(str(value) for value in values if value)
                if source.get("url"):
                    source_text += f"<br/><font color='#65758A' size='7'>{html_escape(str(source['url']))}</font>"
                story.append(Paragraph(f"{source.get('number')}. {source_text}", body))
        story.extend([Spacer(1, 0.5 * cm), Paragraph("本报告基于报告所列资料整理，仅供研究参考，不构成投资建议。", ParagraphStyle("Disclaimer", parent=small, alignment=TA_CENTER))])

        def page(canvas: Any, doc: Any) -> None:
            canvas.saveState()
            canvas.setFont(pdf_font, 8)
            canvas.setFillColor(colors.HexColor("#8A8A8A"))
            if doc.page > 1:
                canvas.drawRightString(A4[0] - 2.25 * cm, A4[1] - 1.25 * cm, f"{prepared['topic']}深度研究报告")
            canvas.drawCentredString(A4[0] / 2, 1.05 * cm, str(doc.page))
            canvas.restoreState()

        document.build(story, onFirstPage=page, onLaterPages=page)
        return buffer.getvalue()

    @classmethod
    def _append_pdf_markdown(cls, story: List[Any], markdown: str, citations: Mapping[str, int], body: Any, h2: Any, h3: Any, Table: Any, TableStyle: Any, colors: Any) -> None:
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.styles import ParagraphStyle
        quote = ParagraphStyle("CNQuote", parent=body, leftIndent=12, borderColor=colors.HexColor(f"#{cls.ACCENT}"), borderWidth=0, borderPadding=7, backColor=colors.HexColor("#F8F2E9"))
        for kind, value in cls._markdown_blocks(markdown):
            if kind == "heading":
                level, text = value
                story.append(__import__("reportlab.platypus", fromlist=["Paragraph"]).Paragraph(html_escape(re.sub(r"^\d+(?:\.\d+)*[\.、\s]+", "", text)), h2 if level <= 2 else h3))
            elif kind == "table":
                rows = [[__import__("reportlab.platypus", fromlist=["Paragraph"]).Paragraph(cls._pdf_text(cell, citations), body) for cell in row] for row in value]
                if rows:
                    table = Table(rows, repeatRows=1, hAlign="CENTER")
                    table.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{cls.TABLE_FILL}")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{cls.ACCENT}")),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C7B9A7")),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]))
                    story.append(table)
            elif kind == "list":
                Paragraph = __import__("reportlab.platypus", fromlist=["Paragraph"]).Paragraph
                for item in value:
                    story.append(Paragraph("• " + cls._pdf_text(item, citations), body))
            elif kind == "quote":
                Paragraph = __import__("reportlab.platypus", fromlist=["Paragraph"]).Paragraph
                story.append(Paragraph(cls._pdf_text(value, citations), quote))
            else:
                Paragraph = __import__("reportlab.platypus", fromlist=["Paragraph"]).Paragraph
                story.append(Paragraph(cls._pdf_text(value, citations), body))

    @staticmethod
    def _pdf_text(value: Any, citations: Mapping[str, int]) -> str:
        text = IndustryResearchReportExportService._clean_text(value, citations)
        escaped = html_escape(text)
        escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
        return re.sub(r"@@CITE:(\d+)@@", r"<super><font color='#9B4B1B'>\1</font></super>", escaped).replace("\n", "<br/>")

    @classmethod
    def _render_figure_png(cls, figure: Mapping[str, Any], path: Path) -> None:
        from PIL import Image, ImageDraw, ImageFont

        width, height = 1400, 760
        image = Image.new("RGB", (width, height), "white")
        draw = ImageDraw.Draw(image)
        font_path = cls._find_cjk_font()
        title_font = ImageFont.truetype(font_path, 42) if font_path else ImageFont.load_default()
        label_font = ImageFont.truetype(font_path, 26) if font_path else ImageFont.load_default()
        small_font = ImageFont.truetype(font_path, 21) if font_path else ImageFont.load_default()
        draw.text((70, 38), str(figure.get("title") or "数据图表"), fill=(45, 45, 45), font=title_font)
        subtitle = str(figure.get("subtitle") or "")
        if subtitle:
            draw.text((70, 92), subtitle, fill=(120, 120, 120), font=small_font)
        left, top, right, bottom = 120, 165, 1320, 650
        draw.line((left, bottom, right, bottom), fill=(190, 190, 190), width=2)
        draw.line((left, top, left, bottom), fill=(190, 190, 190), width=2)
        data = [row for row in figure.get("data") or [] if isinstance(row, Mapping)][:16]
        x_key = str(figure.get("x_key") or figure.get("xKey") or "")
        y_keys = [str(key) for key in (figure.get("y_keys") or figure.get("yKeys") or [])][:4]
        if not data or not y_keys:
            draw.text((left + 50, top + 170), "当前图表数据为空", fill=(130, 130, 130), font=label_font)
            image.save(path, "PNG")
            return
        values = [float(row.get(key) or 0) for row in data for key in y_keys if isinstance(row.get(key), (int, float))]
        max_value = max(values or [1.0]); min_value = min(values or [0.0]); floor = min(0.0, min_value); span = max(1e-9, max_value - floor)
        colors_rgb = [cls.ACCENT_RGB, cls.GOLD_RGB, cls.BLUE_GREY_RGB, (103, 86, 150)]
        chart_type = str(figure.get("type") or "bar")
        group_width = (right - left) / max(1, len(data))
        for row_index, row in enumerate(data):
            center = left + group_width * (row_index + 0.5)
            label = str(row.get(x_key) or row_index + 1)
            if row_index % max(1, len(data) // 8) == 0:
                draw.text((center - 32, bottom + 12), label[:8], fill=(95, 95, 95), font=small_font)
            for key_index, key in enumerate(y_keys):
                raw = row.get(key)
                if not isinstance(raw, (int, float)):
                    continue
                y = bottom - ((float(raw) - floor) / span) * (bottom - top - 25)
                color = colors_rgb[key_index % len(colors_rgb)]
                if chart_type in {"line", "area", "scatter"}:
                    radius = 6
                    draw.ellipse((center - radius, y - radius, center + radius, y + radius), fill=color)
                    if row_index:
                        previous = data[row_index - 1].get(key)
                        if isinstance(previous, (int, float)):
                            previous_y = bottom - ((float(previous) - floor) / span) * (bottom - top - 25)
                            previous_x = left + group_width * (row_index - 0.5)
                            draw.line((previous_x, previous_y, center, y), fill=color, width=4)
                else:
                    bar_width = max(8, group_width * 0.72 / max(1, len(y_keys)))
                    x0 = center - group_width * 0.36 + key_index * bar_width
                    draw.rectangle((x0, y, x0 + bar_width - 3, bottom), fill=color)
        for index, key in enumerate(y_keys):
            x = 80 + index * 250
            draw.rectangle((x, 690, x + 24, 714), fill=colors_rgb[index % len(colors_rgb)])
            draw.text((x + 34, 686), key[:18], fill=(80, 80, 80), font=small_font)
        image.save(path, "PNG")

    @staticmethod
    def _find_cjk_font() -> str | None:
        candidates = (
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        )
        return next((candidate for candidate in candidates if Path(candidate).is_file()), None)
